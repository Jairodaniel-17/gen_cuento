from langchain_openai import ChatOpenAI
from docx import Document
from docx.shared import Cm
from config import LLM_API_KEY, LLM_API_MODEL, LLM_API_URL
from image_generator import ImageGenerator
import time


class StoryGenerator:
    def __init__(self):
        self.llm = ChatOpenAI(
            api_key=LLM_API_KEY, base_url=LLM_API_URL, model=LLM_API_MODEL
        )
        self.image_generator = ImageGenerator()

    def generate_story(self, theme: str, characters: dict, plot: str) -> str:
        """Genera una historia creativa usando OpenAI basada en los parámetros de entrada."""
        characters_list = [
            f"{name} body description: ({info['descripcion_fisica']})"
            for name, info in characters.items()
        ]
        prompt = (
            f"Write a creative story with the theme '{theme}'. Include the characters: {', '.join(characters_list)}. "
            f"The story should follow this plot: {plot}"
            "The story must have at least 6 paragraphs, maximum 10 paragraphs, 4 to 5 lines."
        )
        response = self.llm.invoke(prompt)
        return str(response.content)

    def generate_image(self, description: str) -> str:
        """Genera una imagen usando la API de Stable Diffusion basada en una descripción textual."""
        payload = {
            "prompt": description,
            "negative_prompt": "(an14:1.2), (worst quality, low quality:1.4), (interlocked fingers:1.2), watermark, text",
            "steps": 25,
            "width": 640,
            "height": 640,
            "sampler_index": "Euler a",
        }
        try:
            images = self.image_generator.generate_images([payload])
            return images[0]  # Retorna la ruta de la imagen generada
        except Exception as e:
            print(f"Error generating image: {e}")
            return None

    def save_to_word(
        self, story: str, images: list[str], output_file: str = "story.docx"
    ):
        """Guarda la historia y las imágenes asociadas en un documento Word."""
        doc = Document()
        doc.add_heading("Generated Story", level=1)

        paragraphs = [p.strip() for p in story.split("\n") if p.strip()]

        if len(paragraphs) != len(images):
            print("Warning: The number of paragraphs and images do not match.")

        for idx, paragraph in enumerate(paragraphs):
            # Añade el párrafo al documento y justificar el texto
            doc.add_paragraph(paragraph)
            doc.paragraphs[-1].alignment = 3
            # Si hay una imagen correspondiente, la añade
            if idx < len(images) and images[idx]:
                try:
                    doc.add_picture(images[idx], width=Cm(15), height=Cm(15))
                    doc.paragraphs[-1].alignment = 1  # Centrar la imagen
                except Exception as e:
                    print(f"Failed to add image for paragraph {idx + 1}: {e}")
                    continue

        try:
            doc.save(output_file)
            print(f"Story successfully saved to {output_file}")
        except Exception as e:
            print(f"Failed to save the Word document: {e}")

    def generate_image_descriptions(self, paragraphs, characters: dict) -> list[str]:
        descriptions = []
        for paragraph in paragraphs:
            if not paragraph.strip():
                continue  # Ignorar párrafos vacíos
            instruction = f"""
            Generate a concise description for an image based on this paragraph, also remember to include descriptions of the characters. 
            For women use="1girl", for men use="1boy", and for both use="1girl, 1boy". The descriptions should be descriptive, separated by commas.
            Below is the paragraph and characters to use for generating the description:
            Paragraph: "{paragraph}"
            Characters: {", ".join([f"{name} ({info['descripcion_fisica']})" for name, info in characters.items()])}
            
            **Your description must be in English and include the characters' characteristics, it is for stable diffusion.**
            Description:
            """
            try:
                response = self.llm.invoke(instruction)
                descriptions.append(response.content.strip())
            except Exception as e:
                print(
                    f"Error generating description for paragraph: {paragraph[:50]}... Error: {e}"
                )
                descriptions.append(None)

        return descriptions

    def run(self, theme: str, characters: dict, plot: str):
        # Paso 1: Generar la historia
        print("Generating story...")
        story = self.generate_story(theme, characters, plot)
        print("Story generated successfully.")

        # Paso 2: Generar descripciones de imágenes
        print("Generating image descriptions...")
        paragraphs = [p.strip() for p in story.split("\n") if p.strip()]
        image_descriptions = self.generate_image_descriptions(paragraphs, characters)

        # Paso 3: Generar las imágenes
        print("Generating images...")
        images = [self.generate_image(desc) for desc in image_descriptions if desc]
        print("Images generated successfully.")

        # Paso 4: Guardar en el documento Word
        print("Saving to Word document...")
        timefinish = time.strftime("%Y-%m-%d_%H-%M-%S")
        self.save_to_word(story, images, output_file=f"story_{theme}_{timefinish}.docx")
        print("Process completed.")
